#!/usr/bin/env python3
# -*- coding: utf-8 -*-

"""法律文档向量化与向量数据库存储工具

功能概述：
该工具用于将已标准化和分块哈希的法律文档使用LangChain框架和bge-m3嵌入模型向量化，
并存储到PostgreSQL/pgvector数据库中，以支持高效的语义检索和问答系统。

主要功能：
- 读取已分割的法律条款JSON文件
- 支持.docx文档的直接处理（使用规则进行条款分割）
- 使用bge-m3模型进行文本向量化（支持本地模型检查与下载）
- 连接PostgreSQL/pgvector数据库
- 批量将文档分块和向量存储到数据库
- 提供语义相似度查询功能
- 支持多线程批量处理

使用方法：
1. 确保已安装PostgreSQL和pgvector扩展
2. 配置数据库连接参数
3. 准备已分割的法律条款JSON文件或.docx文档
4. 运行程序，将自动处理文档并存储到向量数据库

输入：已分割的法律条款JSON文件或.docx文档
输出：向量数据存储到PostgreSQL/pgvector数据库
"""

import os
os.environ['TF_CPP_MIN_LOG_LEVEL'] = '2'  # 只显示ERROR和FATAL日志
import re
import json
import time
import hashlib
import threading
import psycopg2
from typing import List, Dict, Optional, Tuple
from concurrent.futures import ThreadPoolExecutor

# 获取脚本所在目录
script_dir = os.path.dirname(os.path.abspath(__file__))

# 抑制TensorFlow警告
import warnings
warnings.filterwarnings('ignore', category=FutureWarning, module='tensorflow')
warnings.filterwarnings('ignore', category=DeprecationWarning, module='tensorflow')

# 抑制LangChain弃用警告
try:
    from langchain_core._api.deprecation import LangChainPendingDeprecationWarning
    warnings.filterwarnings('ignore', category=LangChainPendingDeprecationWarning)
except ImportError:
    # 如果无法导入警告类，忽略这个错误
    pass

# LangChain相关导入
# 强制使用新版本的HuggingFace嵌入
try:
    from langchain_huggingface import HuggingFaceEmbeddings
    USE_NEW_EMBEDDINGS = True
except ImportError as e:
    raise ImportError("请安装 langchain-huggingface 包: pip install langchain-huggingface") from e

from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.document_loaders import JSONLoader, TextLoader

# 强制使用旧的 langchain_community 实现（避免依赖问题）
from langchain_community.vectorstores import PGVector
USE_NEW_PGVECTOR = False

from langchain.schema import Document

# Transformers相关导入（用于模型下载）
from transformers import AutoTokenizer, AutoModel

# 文档处理相关导入
import docx

# 用于模型下载的导入
from transformers import AutoTokenizer, AutoModel

# 获取脚本所在目录
script_dir = os.path.dirname(os.path.abspath(__file__))

class LegalVectorDBManager:
    """法律文档向量化与向量数据库管理类
    
    提供法律文档读取、分块、向量化和向量数据库存储功能
    """
    
    def __init__(self, db_connection_string: str, collection_name: str = "legal_documents"):
        """初始化法律文档向量化与向量数据库管理器
        
        Args:
            db_connection_string: PostgreSQL数据库连接字符串
            collection_name: 向量集合名称
        """
        # 数据库配置
        self.db_connection_string = db_connection_string
        self.collection_name = collection_name
        
        # 模型相关配置
        self.model_name = "BAAI/bge-m3"
        # 优先使用用户缓存目录中的模型
        user_cache_dir = os.path.expanduser("~/.cache/huggingface/hub")
        # Hugging Face缓存目录结构：models--{org}--{model_name}/snapshots/{commit_hash}
        model_cache_dir = os.path.join(user_cache_dir, "models--BAAI--bge-m3")
        
        # 查找最新的快照目录
        if os.path.exists(model_cache_dir):
            snapshots_dir = os.path.join(model_cache_dir, "snapshots")
            if os.path.exists(snapshots_dir):
                # 获取最新的快照目录
                snapshot_dirs = [d for d in os.listdir(snapshots_dir) if os.path.isdir(os.path.join(snapshots_dir, d))]
                if snapshot_dirs:
                    # 使用最新的快照（按修改时间排序）
                    latest_snapshot = max(snapshot_dirs, key=lambda x: os.path.getmtime(os.path.join(snapshots_dir, x)))
                    self.local_model_dir = os.path.join(snapshots_dir, latest_snapshot)
                else:
                    self.local_model_dir = model_cache_dir
            else:
                self.local_model_dir = model_cache_dir
        else:
            # 如果缓存目录不存在，直接使用模型名称让HuggingFace自动处理
            self.local_model_dir = self.model_name
        # 备用路径：项目本地模型目录
        self.fallback_model_dir = os.path.join(os.path.dirname(os.getcwd()), "..\11_local_models", "bge-m3")
        
        # 检查本地是否有模型文件，如果没有则下载
        self._check_and_download_model()
        
        # 初始化嵌入模型 - 使用本地bge-m3模型
        # 强制使用新版本的HuggingFace嵌入
        self.embeddings = HuggingFaceEmbeddings(
            model_name=self.local_model_dir,  # 使用本地模型目录
            model_kwargs={"device": "cpu"},  # 可根据实际情况改为"cuda"
            encode_kwargs={"normalize_embeddings": True}
        )
        
        # 抑制sentence-transformers的警告
        import logging
        logging.getLogger("sentence_transformers").setLevel(logging.ERROR)
        
        # 初始化文本分割器
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=500,
            chunk_overlap=50,
            separators=["。", "；", "！", "？", "\n", "\r", " "]
        )
        
        # 法律文本特征模式（用于规则分割）
        self.article_patterns = [
            r'第\s*[零一二三四五六七八九十百千]+\s*条',  # 中文数字条
            r'第\s*\d+\s*条',  # 阿拉伯数字条
            r'凡\s*[^。；；]+者',  # 凡...者句式
            r'[一二三四五六七八九十]+、',  # 项号标记
            r'第\s*[一二三四五六七八九十百千]+\s*款',  # 款号标记
        ]
        
        # 编译正则表达式以提高效率
        self.compiled_patterns = [re.compile(pattern) for pattern in self.article_patterns]
        
        # 向量数据库实例
        self.vector_store = None
        
        # 初始化数据库连接
        self._init_vector_store()
    
    def _check_and_download_model(self):
        """检查本地是否有模型文件，如果没有则下载"""
        # 检查本地模型目录是否存在且包含必要文件
        required_files = ["config.json", "tokenizer.json", "model.safetensors"]
        model_exists = False
        
        # 首先检查用户缓存目录
        if os.path.exists(self.local_model_dir):
            # 检查是否包含必要的模型文件
            model_exists = True
            for file in required_files:
                if not os.path.exists(os.path.join(self.local_model_dir, file)):
                    model_exists = False
                    break
        
        # 如果用户缓存目录没有，检查备用目录
        if not model_exists and os.path.exists(self.fallback_model_dir):
            model_exists = True
            for file in required_files:
                if not os.path.exists(os.path.join(self.fallback_model_dir, file)):
                    model_exists = False
                    break
            # 如果备用目录有模型，使用备用目录
            if model_exists:
                self.local_model_dir = self.fallback_model_dir
        
        if not model_exists:
            print(f"本地未找到完整的bge-m3模型，将使用HuggingFace自动下载机制")
            print("模型将自动下载到用户缓存目录：~/.cache/huggingface/hub/")
            # 直接使用模型名称，让HuggingFaceBgeEmbeddings自动处理下载和缓存
            self.local_model_dir = self.model_name
        else:
            print(f"找到本地bge-m3模型，将使用本地模型进行向量化")
    
    def _init_vector_store(self):
        """初始化向量数据库连接"""
        try:
            print("正在连接PostgreSQL/pgvector数据库...")
            
            if USE_NEW_PGVECTOR:
                # 使用新的 langchain_postgres 实现
                # 新版本使用 embedding_function 参数而不是 embedding
                self.vector_store = PGVector(
                    connection_string=self.db_connection_string,
                    collection_name=self.collection_name,
                    embedding_function=self.embeddings,
                    use_jsonb=True  # 使用JSONB格式避免警告
                )
            else:
                # 使用旧的 langchain_community 实现
                # 旧版本也使用 embedding_function 参数
                self.vector_store = PGVector(
                    connection_string=self.db_connection_string,
                    collection_name=self.collection_name,
                    embedding_function=self.embeddings,
                    use_jsonb=True  # 使用JSONB格式避免警告
                )
            
            print("数据库连接成功！")
        except Exception as e:
            print(f"数据库连接失败: {str(e)}")
            print("请确保已安装PostgreSQL和pgvector扩展，并正确配置连接字符串")
            # 即使连接失败也继续执行，后面会检查vector_store是否为None
            self.vector_store = None
    
    def _generate_sha256(self, text: str) -> str:
        """为文本生成SHA-256哈希值"""
        return hashlib.sha256(text.encode('utf-8')).hexdigest()
    
    def read_docx_file(self, file_path: str) -> str:
        """读取Word文档内容"""
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"文件不存在: {file_path}")
        
        doc = docx.Document(file_path)
        full_text = []
        
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:  # 跳过空段落
                full_text.append(text)
        
        return "\n".join(full_text)
    
    def detect_article_boundaries(self, text: str) -> List[int]:
        """检测条款边界位置"""
        boundaries = [0]  # 起始位置
        
        # 查找所有匹配的模式位置
        for pattern in self.compiled_patterns:
            for match in pattern.finditer(text):
                # 如果是新的边界位置且不在已有边界中
                if match.start() not in boundaries:
                    boundaries.append(match.start())
        
        # 添加文本结束位置
        boundaries.append(len(text))
        
        # 排序边界位置
        boundaries.sort()
        
        return boundaries
    
    def split_by_rules(self, text: str) -> List[str]:
        """使用规则进行文本分割"""
        chunks = []
        boundaries = self.detect_article_boundaries(text)
        
        # 根据边界分割文本
        for i in range(len(boundaries) - 1):
            start, end = boundaries[i], boundaries[i + 1]
            chunk = text[start:end].strip()
            if chunk:  # 跳过空块
                chunks.append(chunk)
        
        return chunks
    
    def segment_document(self, file_path: str) -> List[Dict]:
        """分割文档并返回结构化的条款列表"""
        print(f"正在处理文档: {file_path}")
        text = self.read_docx_file(file_path)
        
        # 使用规则进行分割
        articles = self.split_by_rules(text)
        
        # 结构化每个条款
        structured_articles = []
        for i, article_text in enumerate(articles, 1):
            article_info = {
                'id': str(i),
                'text': article_text,
                'length': len(article_text),
                'has_article_number': any(pattern.search(article_text) for pattern in self.compiled_patterns),
                'start_with': article_text[:100] + '...' if len(article_text) > 100 else article_text
            }
            structured_articles.append(article_info)
        
        print(f"文档分割完成，共提取 {len(structured_articles)} 条法律条款")
        return structured_articles
    
    def load_articles_from_json(self, file_path: str) -> List[Dict]:
        """从JSON文件加载已分割的法律条款"""
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"文件不存在: {file_path}")
        
        print(f"正在加载条款文件: {file_path}")
        with open(file_path, 'r', encoding='utf-8') as f:
            data = json.load(f)
        
        # 处理不同类型的JSON格式
        if isinstance(data, list):
            # 如果是列表格式，直接返回
            articles = data
            print(f"成功加载 {len(articles)} 条法律条款")
            return articles
        elif isinstance(data, dict):
            # 如果是字典格式，检查是否包含条款数据
            articles = []
            for key, value in data.items():
                if isinstance(value, dict):
                    # 如果值是字典，检查是否包含文本内容
                    if 'text' in value:
                        article = value.copy()
                        article['id'] = key
                        articles.append(article)
                    elif 'content' in value:
                        article = {'id': key, 'text': value['content']}
                        if 'hash' in value:
                            article['hash'] = value['hash']
                        articles.append(article)
                else:
                    # 如果值不是字典，可能是简单的文本
                    articles.append({'id': key, 'text': str(value)})
            
            if articles:
                print(f"成功从字典格式加载 {len(articles)} 条法律条款")
                return articles
            else:
                print("警告: JSON文件是字典格式但不包含可识别的条款数据")
                return []
        else:
            print(f"警告: 未知的JSON格式: {type(data)}")
            return []
    
    def convert_to_documents(self, articles: List[Dict], source_file: str) -> List[Document]:
        """将条款列表转换为LangChain的Document对象列表"""
        documents = []
        
        for article in articles:
            # 提取元数据
            metadata = {
                'source': source_file,
                'article_id': article.get('id', 'unknown'),
                'length': article.get('length', len(article.get('text', ''))),
                'has_article_number': article.get('has_article_number', False),
                'import_time': time.strftime('%Y-%m-%d %H:%M:%S')
            }
            
            # 如果有哈希值，添加到元数据
            if 'hash' in article:
                metadata['hash'] = article['hash']
            else:
                # 为文本生成哈希值
                text = article.get('text', '')
                metadata['hash'] = self._generate_sha256(text)
            
            # 创建Document对象
            document = Document(
                page_content=article.get('text', ''),
                metadata=metadata
            )
            documents.append(document)
        
        return documents
    
    def process_single_file(self, file_path: str) -> bool:
        """处理单个文件，将其向量化并存储到数据库"""
        try:
            # 检查文件类型并加载条款
            if file_path.endswith('.json'):
                # 假设JSON文件是已分割的条款文件
                articles = self.load_articles_from_json(file_path)
            elif file_path.endswith('.docx'):
                # 对于Word文档，需要先进行分割
                articles = self.segment_document(file_path)
            else:
                print(f"不支持的文件类型: {file_path}")
                return False
            
            # 转换为Document对象
            documents = self.convert_to_documents(articles, file_path)
            
            # 进一步分割长文本（如果需要）
            split_documents = []
            for doc in documents:
                if len(doc.page_content) > 500:
                    # 对于过长的文本，进行进一步分割
                    splits = self.text_splitter.split_documents([doc])
                    split_documents.extend(splits)
                else:
                    split_documents.append(doc)
            
            # 向量化并存储到数据库
            if self.vector_store:
                print(f"正在向量化并存储 {len(split_documents)} 个文档块...")
                start_time = time.time()
                
                # 批量添加文档到向量数据库
                self.vector_store.add_documents(split_documents)
                
                end_time = time.time()
                print(f"文档存储完成，耗时: {end_time - start_time:.2f} 秒")
                return True
            else:
                print("向量数据库未初始化，无法存储文档")
                return False
        except Exception as e:
            print(f"处理文件 {file_path} 时发生错误: {str(e)}")
            return False
    
    def batch_process_files(self, directory_path: str, max_workers: int = 4) -> Tuple[int, int]:
        """批量处理目录中的所有文件
        
        Args:
            directory_path: 要处理的目录路径
            max_workers: 最大工作线程数
        
        Returns:
            Tuple[int, int]: (成功处理的文件数, 总文件数)
        """
        if not os.path.exists(directory_path):
            print(f"目录不存在: {directory_path}")
            return 0, 0
        
        # 获取目录中的所有支持的文件
        supported_files = []
        for root, _, files in os.walk(directory_path):
            for file in files:
                if file.endswith('.json') or file.endswith('.docx'):
                    supported_files.append(os.path.join(root, file))
        
        total_files = len(supported_files)
        success_count = 0
        
        if total_files == 0:
            print(f"目录 {directory_path} 中没有找到支持的文件")
            return 0, 0
        
        print(f"找到 {total_files} 个支持的文件，开始批量处理...")
        
        # 使用线程池进行并行处理
        if total_files > 1 and max_workers > 1:
            with ThreadPoolExecutor(max_workers=min(max_workers, total_files)) as executor:
                # 提交所有任务
                futures = {executor.submit(self.process_single_file, file_path): file_path for file_path in supported_files}
                
                # 收集结果
                for future in futures:
                    try:
                        if future.result():
                            success_count += 1
                    except Exception as e:
                        print(f"处理文件 {futures[future]} 时发生异常: {str(e)}")
        else:
            # 单线程处理
            for file_path in supported_files:
                if self.process_single_file(file_path):
                    success_count += 1
        
        print(f"批量处理完成，成功处理 {success_count}/{total_files} 个文件")
        return success_count, total_files
    
    def similarity_search(self, query: str, k: int = 5) -> List[Dict]:
        """执行语义相似度搜索
        
        Args:
            query: 搜索查询文本
            k: 返回的最相似结果数量
        
        Returns:
            List[Dict]: 搜索结果列表
        """
        if not self.vector_store:
            print("向量数据库未初始化，无法执行搜索")
            return []
        
        try:
            print(f"正在执行相似度搜索: {query}")
            results = self.vector_store.similarity_search_with_score(query, k=k)
            
            # 格式化搜索结果
            formatted_results = []
            for doc, score in results:
                result = {
                    'text': doc.page_content,
                    'metadata': doc.metadata,
                    'similarity_score': 1 - score  # 转换为相似度得分（越高越相似）
                }
                formatted_results.append(result)
            
            return formatted_results
        except Exception as e:
            print(f"执行搜索时发生错误: {str(e)}")
            return []
    
    def count_documents(self) -> int:
        """获取数据库中的文档数量"""
        if not self.vector_store:
            print("向量数据库未初始化")
            return 0
        
        try:
            # 使用psycopg2直接查询数据库获取文档数量
            conn = psycopg2.connect(self.db_connection_string)
            cursor = conn.cursor()
            
            # 查询指定集合中的文档数量
            cursor.execute("""SELECT COUNT(*) FROM langchain_pg_embedding WHERE collection_id = (
                SELECT id FROM langchain_pg_collection WHERE name = %s
            )""", (self.collection_name,))
            
            count = cursor.fetchone()[0]
            
            cursor.close()
            conn.close()
            
            return count
        except Exception as e:
            print(f"获取文档数量时发生错误: {str(e)}")
            return 0


def check_pgvector_installation(db_params):
    """检查pgvector扩展是否已安装"""
    try:
        import psycopg2
        # 连接到数据库
        conn = psycopg2.connect(
            host=db_params["host"],
            port=db_params["port"],
            dbname=db_params["dbname"],
            user=db_params["user"],
            password=db_params["password"]
        )
        cursor = conn.cursor()
        
        # 检查是否已安装vector扩展
        cursor.execute("SELECT EXISTS(SELECT 1 FROM pg_extension WHERE extname = 'vector')")
        vector_installed = cursor.fetchone()[0]
        
        cursor.close()
        conn.close()
        
        if vector_installed:
            print("✓ pgvector扩展已安装")
            return True
        else:
            print("✗ pgvector扩展未安装")
            print("请手动安装pgvector扩展：")
            print("1. 下载pgvector: https://github.com/pgvector/pgvector")
            print("2. 编译安装（Windows用户可以使用预编译版本）")
            print("3. 将vector.dll和vector.control文件复制到PostgreSQL的extension目录")
            print("4. 在PostgreSQL中执行: CREATE EXTENSION vector")
            print("或者使用以下Docker命令快速启动带pgvector的PostgreSQL:")
            print("docker run -d --name postgres-vector -p 5432:5432 -e POSTGRES_PASSWORD=your_password ankane/pgvector")
            return False
        
    except Exception as e:
        print(f"检查pgvector扩展时发生错误: {str(e)}")
        print("请确保PostgreSQL服务正在运行且连接参数正确")
        return False

def main():
    """主函数"""
    # 配置数据库连接参数
    DB_PARAMS = {
        "host": "localhost",
        "port": 5432,
        "dbname": "legal_database",
        "user": "postgres",
        "password": "pgadmin"
    }
    
    # 首先检查pgvector扩展是否已安装
    if not check_pgvector_installation(DB_PARAMS):
        print("无法继续，请先安装pgvector扩展")
        print("或者您可以选择使用其他向量数据库（如Chroma、FAISS等）")
        return
    
    # 构建数据库连接字符串
    db_connection_string = f"postgresql://{DB_PARAMS['user']}:{DB_PARAMS['password']}@{DB_PARAMS['host']}:{DB_PARAMS['port']}/{DB_PARAMS['dbname']}"
    
    # 创建法律文档向量化与向量数据库管理器实例
    try:
        vector_db_manager = LegalVectorDBManager(db_connection_string)
        
        # 直接使用当前目录
        target_directory = "."
        print("使用当前目录进行文件处理")
        
        # 执行批量处理
        success_count, total_files = vector_db_manager.batch_process_files(target_directory)

        
        # 统计数据库中的文档数量
        doc_count = vector_db_manager.count_documents()
        print(f"向量数据库中共有 {doc_count} 个文档块")
        
        # 示例：执行相似度搜索
        # query = "合同纠纷如何解决？"
        # search_results = vector_db_manager.similarity_search(query, k=3)
        # print(f"\n搜索结果 ({len(search_results)}):")
        # for i, result in enumerate(search_results, 1):
        #     print(f"\n结果 {i} (相似度: {result['similarity_score']:.4f}):")
        #     print(f"文本: {result['text'][:200]}...")
        #     print(f"来源: {result['metadata'].get('source', '未知')}")
            
    except Exception as e:
        print(f"程序运行失败: {str(e)}")


if __name__ == "__main__":
    main()